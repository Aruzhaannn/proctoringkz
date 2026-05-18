# -*- coding: utf-8 -*-
from docx import Document

doc = Document()
doc.add_heading('Емтихан Сұрақтары', 0)

questions = [
    ('Java бағдарламалау тілінің авторы кім?', ['Джеймс Гослинг', 'Билл Гейтс', 'Стив Джобс', 'Линус Торвальдс'], 'A'),
    ('SQL сөзі қалай жіктеледі?', ['Structured Query Language', 'Simple Question Language', 'Strong Query Logic', 'System Query Language'], 'A'),
    ('HTTP портының стандартты нөмірі қандай?', ['80', '443', '21', '22'], 'A'),
    ('Git-те жаңа тармақ (branch) құру бұйрығы қандай?', ['git branch', 'git checkout', 'git commit', 'git pull'], 'A'),
    ('Python-да тізімге элемент қосу үшін қандай әдіс қолданылады?', ['append()', 'add()', 'insert()', 'push()'], 'A'),
    ('HTML-дің толық атауы қандай?', ['HyperText Markup Language', 'Hyper Tool Multi Language', 'High Text Machine Language', 'Hyperlinks Text Mark Language'], 'A'),
    ('CSS-тің мақсаты неде?', ['Бетті безендіру', 'Серверлік логика', 'Деректер қорын басқару', 'Желілік трафикті реттеу'], 'A'),
    ('Docker-дің негізгі мақсаты қандай?', ['Контейнеризация', 'Виртуализация', 'Компиляция', 'Шифрлау'], 'A'),
    ('Spring Boot қай тілде жазылған?', ['Java', 'C++', 'Python', 'Go'], 'A'),
    ('React кітапханасын қай компания жасады?', ['Facebook', 'Google', 'Microsoft', 'Amazon'], 'A'),
    ('REST API қандай форматты жиі қолданады?', ['JSON', 'XML', 'YAML', 'CSV'], 'A'),
    ('PostgreSQL қандай дерекқор түріне жатады?', ['Реляциялық', 'NoSQL', 'Graph', 'Document-oriented'], 'A'),
    ('Linux жүйесінде файлды жою үшін қандай бұйрық қолданылады?', ['rm', 'del', 'erase', 'remove'], 'A'),
    ('AWS қандай қызметтер ұсынады?', ['Бұлтты инфрақұрылым', 'Автомобиль өндірісі', 'Ойындар жасау', 'Азық-түлік жеткізу'], 'A'),
    ('Бағдарламалық жасақтаманы тестілеу (QA) не үшін керек?', ['Қателерді табу және сапаны арттыру', 'Кодты тез жазу', 'Серверді жылдамдату', 'Қолданушыларды бақылау'], 'A')
]

for i, (q, opts, ans) in enumerate(questions, 1):
    doc.add_paragraph(f'{i}. {q}')
    doc.add_paragraph(f'A) {opts[0]}')
    doc.add_paragraph(f'B) {opts[1]}')
    doc.add_paragraph(f'C) {opts[2]}')
    doc.add_paragraph(f'D) {opts[3]}')
    doc.add_paragraph(f'Answer: {ans}')
    doc.add_paragraph('')

doc.save('exam_15_questions.docx')
print('exam_15_questions.docx created successfully!')
